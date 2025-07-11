from Bio import SeqIO
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Cm
from docx.shared import Pt
from methpy.__init__ import find_start, Gene_name
import os
import sys
import tkinter as tk
from tkinter import ttk


def check ():

    """Checks a sequence of DNA treated with bisulfite assay, both in CpG and non CpG sites"""
    #style of the word document
    doc = Document ()
    style = doc.styles ["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = Pt (13)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm (2.5)
        section.bottom_margin = Cm (2.5)
        section.left_margin = Cm (2.5)
        section.right_margin = Cm (2.5)

    #takes the informations and then split them
    informations = Gene_name ().response
    
    #checks if there are all the needed informations
    #tries to open files, otherwise asks for valid info
    try:
        gene_name_str = informations ["reference"]
        seq_file_path = informations ["file"]
        script_path = os.getcwd ()
        path_forward = os.path.join (script_path,"References", gene_name_str + "F.txt")
        path_reverse = os.path.join (script_path, "References", gene_name_str + "R.txt")

        #opens both the reference forward and reverse
        forward = open (path_forward, "r")
        forward = forward.read ()
        reverse = open (path_reverse, "r")
        reverse = reverse.read ()

        #takes the sequence name
        name_sequence = os.path.basename (seq_file_path)
        name_sequence = os.path.splitext (name_sequence) [0]

        #creates paragraphs on word
        p = doc.add_paragraph (name_sequence)
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 0
        p.style = doc.styles ["Normal"]
        p.add_run ("")

        #opens the sequence
        seq_file_path = seq_file_path.replace ('"','')
        seq = open (seq_file_path)
        if os.path.splitext (seq_file_path)[-1] == ".ab1":
            seq = SeqIO.read(seq_file_path, "abi").seq
        else: seq = seq.read ()
        seq = seq.replace ("\n", "")
        
        #checks if which one between the reference and the sequence is longer
        if len(seq)  <len(forward):
            number = range (len(seq))
        else:
            number = range (len(forward))

        #takes informations about the start
        start_informations = find_start(forward, reverse, seq, number)

        if start_informations != "Start not found" and start_informations != None:
            #takes the informations and splits them
            start = int (start_informations[0])
            start_reference= int (start_informations[1])
            reference = str (start_informations[2])
            
            #If there’s a part of my sequence before the start, I can see those letters, even if there are some mistakes
            if start > start_reference:
                seq = seq [start-start_reference:]
            #otherwise it inserts Xs
            else:
                seq = "X" * (start_reference-start) + seq

            class PopupChoice:  

                def __init__ (self, skip_popup):
                    self.response = {"response":"", "choice":"", "index":"", "doubtful_index":"", "stop":""}

                    #if the error is an deletion, adds the base
                    def deletion ():
                        newseq = seq [0:(n)] + ref [n] + seq [(n):] 
                        self.response ["response"] = newseq
                        #takes into account if the insertion is about a cytosine
                        if (ref [n] == "C" and ref == forward) or (ref[n] == "G" and ref == reverse):
                            self.response ["doubtful_index"] = n
                            #adds the letter into the word file and highlights it
                            p.add_run (
                                ref[n]
                            ).font.highlight_color = WD_COLOR_INDEX.RED
                            self.response ["choice"] = "doubtful_deletion"
                        else:
                            p.add_run (
                                ref[n]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                            self.response ["choice"] = "deletion"
                        popup.quit ()

                    #if the error is a insertion, deletes the base present in the reference
                    def insertion ():
                        newseq = seq [0:(n)] + seq [(n+1):]
                        self.response ["index"] = n
                        self.response ["response"] = newseq
                        self.response ["choice"] = "insertion"
                        popup.quit ()

                    #in case of a base change, highlights it in the word document
                    def base_change ():
                        newseq = seq
                        self.response["response"] = newseq
                        #distinguishes between "doubtful" or not
                        if (ref [n] == "C" and ref == forward) or (ref[n] == "G" and ref == reverse):
                            self.response ["doubtful_index"] = n
                            p.add_run ( 
                                seq[n]
                            ).font.highlight_color = WD_COLOR_INDEX.BLUE
                            self.response ["choice"] = "doubtful_base_change"
                        else:
                            p.add_run (
                                seq[n]
                            ).font.highlight_color = WD_COLOR_INDEX.PINK
                            self.response ["choice"] = "choice_ base_change"
                        popup.quit ()

                    #this to exit the loop
                    def stop (): 
                        self.response ["stop"] = "stop"
                        popup.quit ()
                        exit
                    
                    if not skip_popup:
                        #the popup characteristics
                        popup = tk.Tk ()
                        if sys.platform.startswith ('darwin'): 
                            font = ("Courier New", 19)
                        else: font = ("Courier New", 13)
                        popup.geometry ("600x200")
                        popup.wm_title (title)
                        label = ttk.Label (popup, text = msg, font = font)
                        label.place (rely = 0.25, relx = 0.5, anchor = "center")  
                        style_b1 = ttk.Style()
                        style_b1.configure ("B1.TButton", foreground = "#008b8b")
                        b1 = ttk.Button (popup, text = "Deletion", command = deletion, style = "B1.TButton")
                        style_b2 = ttk.Style()
                        style_b2.configure ("B2.TButton", foreground = "magenta")
                        b2 = ttk.Button (popup, text = "Base Change", command = base_change, style = "B2.TButton")
                        style_b3 = ttk.Style()
                        style_b3.configure ("B3.TButton", foreground = "red")                                
                        b3 = ttk.Button (popup, text = "Stop", command = stop, style = "B3.TButton")
                        b4 = ttk.Button (popup, text = "Insertion", command = insertion)
                        b1.place(rely = 0.5,relx = 0.25, anchor = "center") 
                        b4.place (rely = 0.5,relx = 0.5, anchor = "center")
                        b2.place (rely = 0.5,relx = 0.75, anchor = "center")
                        b3.place (rely = 0.75, relx = 0.5, anchor = "center")

                        popup.mainloop ()
                        popup.destroy ()

            methylated_indices = ""
            removed_indices = ""
            doubtful_indices = ""

            i = 0
            msg = ""
            title = "First base is an error, select what type:"

            if reference == forward:
                p = doc.add_paragraph ("Forward" + "\n")
            elif reference == reverse:
                p = doc.add_paragraph ("Reverse" + "\n")
            
            skip_space = False

            while i < (len(seq)) and i < (len(reference)): 
                skip_popup = True
                output_choice = {"response":"", "choice":"", "stop":""}
                
                if skip_space == False:
                    #inserts number at the beginning of a new line 
                    if i % 50 == 0 and i > 1: 
                        if reference == forward:
                            base_number = 1+i
                            base_number = str (base_number)
                        #if the reference is a reverse, it starts counting from the last one and going backwards
                        else: 
                            base_number = len (reference)
                            base_number = base_number - i
                            base_number = str (base_number)

                        #adds into the word files both spaces and numbers
                        p.add_run (" ") 
                        p.add_run (base_number)
                        p.add_run (" ")       

                        if int (base_number) < 100:
                            p.add_run (" ")

                        if int (base_number) < 10: 
                            p.add_run (" ")

                    #inserts the correct number of spaces before the first nucleotide
                    elif i == 0:
                        if reference == forward:
                            base_number = 1
                            base_number = str (base_number)
                            p.add_run (base_number)
                            p.add_run ("   ")
                        else:
                            base_number = len (reference)
                            base_number = str (base_number) 
                            p.add_run (base_number)
                            p.add_run (" ")   
                            if len (reference) < 100 :
                                p.add_run (" ")

                    #inserts a space every 10nt
                    elif (i) % 10 == 0 and i > 1 :
                        p.add_run (" ")  
    
                skip_space = False
                #inserts Xs at the beginning, if the sequence doesn't have the first nucleotides
                if seq [i] == "X":
                    if (reference [i] == "C" and reference == forward) or (reference[i] == "G" and reference == reverse):
                        p.add_run (
                            seq[i]
                        ).font.highlight_color = WD_COLOR_INDEX.RED

                        if doubtful_indices != "":
                            doubtful_indices = doubtful_indices + ", "

                        if reference == forward:
                            doubtful_indices = doubtful_indices + str (i+1) 
                        elif reference == reverse :
                            doubtful_indices = doubtful_indices + str (len(reference)-i) 

                    else:
                        p.add_run(
                                seq[i]
                            ).font.highlight_color = WD_COLOR_INDEX.TURQUOISE

                #if it is a methylated position, it is highlighted in green       
                elif ((seq[i] == "C" and reference [i] == "C") and reference == forward) or (seq[i] == "G" and reference[i] == "G" and reference == reverse):
                    p.add_run (
                        seq[i]
                    ).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

                    if reference == forward:
                        base_number == str (i+1)
                    elif reference == reverse:
                        base_number = str (len(reference)-i-1)

                    if methylated_indices != "":
                        methylated_indices = methylated_indices + ", "

                    if reference == forward:
                            methylated_indices = methylated_indices + str (i+1)
                    elif reference == reverse:
                            methylated_indices = methylated_indices + str (len(reference)-i)

                #in case that base is not methylated, it is highlighted in yellow
                elif ((seq [i] == "T" and reference[i] == "C") and reference == forward) or (seq[i] == "A" and reference[i] == "G" and reference == reverse):
                    p.add_run (
                        seq[i]
                    ).font.highlight_color = WD_COLOR_INDEX.YELLOW

                #in case the sequence is not equal to the reference, it will show the popup
                elif seq [i] != reference [i]:
                    if reference == forward:
                        msg = "SEQUENCE "+ "\t" + seq [i:(i+20)] + "\n" + "FORWARD  " + "\t" + reference [i:(i+20)]
                    elif reference == reverse:
                        msg = "SEQUENCE "+ "\t" + seq [i:(i+20)] + "\n" + "REVERSE  " + "\t" + reference [i:(i+20)]

                    skip_popup = False
                    n = i 
                    ref = reference
                    output_choice = PopupChoice (skip_popup).response
                    seq = output_choice ["response"]
                else:
                    p.add_run (seq[i])

                if output_choice ["choice"] == "insertion":
                    if removed_indices != "":
                        removed_indices =  removed_indices + ", " 
                    if reference == forward:
                        removed_indices = removed_indices + str (i+1)
                    elif reference == reverse:
                        removed_indices = removed_indices + str (len(reference)-i)
    
                    skip_space = True
                    next 

                elif output_choice ["choice"] == "doubtful_base_change" or output_choice ["choice"] == "doubtful_deletion":
                    if doubtful_indices != "":
                        doubtful_indices = doubtful_indices + ", "
                    if reference == forward:
                        doubtful_indices = doubtful_indices + str (i+1) 
                    elif reference == reverse :
                        doubtful_indices=doubtful_indices + str (len(reference)-i) 

                    i += 1
                    next
                    
                else:
                    i += 1
                

            #adds the legend and all the info at the end                
            p = doc.add_paragraph () 
            p.add_run ("\n" + "Legend:" + "\n")
            p.add_run ("Methylated" + "\n").font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            p.add_run ("Non-Methylated" + "\n").font.highlight_color = WD_COLOR_INDEX.YELLOW
            p.add_run ("Base Change" + "\n").font.highlight_color = WD_COLOR_INDEX.PINK
            p.add_run ("Doubtful Base Change" + "\n").font.highlight_color = WD_COLOR_INDEX.BLUE
            p.add_run ("Deletion" + "\n").font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            p.add_run ("Doubtful Deletion" + "\n").font.highlight_color = WD_COLOR_INDEX.RED
            p.add_run ("\n" + "\n" + "Insertion positions:" + "\n" + removed_indices)
            p.add_run ("\n" + "\n" + "Methylated C positions:"+ "\n" + methylated_indices)
            p.add_run ("\n" + "\n" + "Doubtful C positions:"+ "\n" + doubtful_indices)

            #saves the file as a word document in "Output in word"
            word_file_path = os.path.join (os.path.splitext (seq_file_path)[0] + ".docx")
            output_word_file_path = word_file_path.replace ("Input", "Output in word")
            output_word_without_name = os.path.dirname (os.path.abspath(output_word_file_path))

            if os.path.isdir (output_word_without_name) == False:
                os.makedirs (output_word_without_name)

            try: doc.save(output_word_file_path)
            except: print ("Failed to save the Word file")

            #saves the file as a .txt in "Output in txt"
            output_txt_without_name = output_word_without_name.replace ("Output in word", "Output in txt")
            output_txt_file_path = output_word_file_path.replace ("Output in word", "Output in txt").replace ("docx", "txt")
            if os.path.isdir (output_txt_without_name) == False: 
                os.makedirs (output_txt_without_name)

            #creates the content of the txt file
            if reference == forward:
                if output_choice ["stop"] == "stop":
                    final_string = name_sequence + "\n" + methylated_indices.replace (", ", "\n") + "\n" + "doubtful_c_positions" + "\n" + doubtful_indices.replace (", ", "\n") + "\n"+ str (i-1) + "\n" + "F"
                else: 
                    final_string = name_sequence + "\n" + methylated_indices.replace (", ", "\n") + "\n" + "doubtful_c_positions" + "\n" + doubtful_indices.replace (", ", "\n") + "\n"+ str (i) + "\n" + "F"
            else:  
                if output_choice ["stop"] == "stop":
                    final_string = name_sequence + "\n" + methylated_indices.replace (", ", "\n") + "\n" + "doubtful_c_positions" + "\n" + doubtful_indices.replace (", ", "\n") + "\n"+ str (len(reference)-i+1+1) + "\n"+ "R"
                else:
                    final_string = name_sequence + "\n" + methylated_indices.replace (", ", "\n") + "\n" + "doubtful_c_positions" + "\n" + doubtful_indices.replace (", ", "\n") + "\n"+ str (len(reference)-i+1) + "\n"+ "R"

            with open (output_txt_file_path , "w") as f:
                f.write (final_string)

            folder_file = seq_file_path.split ("Input") [1]
            file_name = os.path.basename (folder_file)
            folder_name = folder_file.split (file_name) [0]
            print ("Just checked", file_name, "in", folder_name[1:-1])
            
        else: print ("Check failed to find the start.")

    except:
        print ("You sequence was not found, try selecting a valid file or a valid Gene name.")

sys.modules[__name__] = check
